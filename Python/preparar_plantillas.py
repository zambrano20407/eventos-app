"""
preparar_plantillas.py
Convierte los formatos Word oficiales de "Formato Usuarios" en plantillas
con marcadores, listas para rellenar desde el navegador.

Trabaja con python-docx (entiende tablas, filas y celdas), NO con regex
sobre el XML: asi el documento conserva intactos formato y fuentes.

Que hace:
  1. {{campo}} en la celda vacia que sigue a cada etiqueta conocida.
  2. {{cb:opcion}} en lugar del caracter de cada casilla de verificacion.
  3. Imagen de firma (placeholder PNG) en la celda "Firma:" del funcionario;
     el navegador solo reemplaza ese binario por la firma real.

Ejecutar cuando cambien los formatos oficiales:
    py Python/preparar_plantillas.py
"""
import os
import re
import shutil
import zipfile

import docx
from docx.shared import Cm, Pt

PYTHON_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(PYTHON_DIR)
ORIGEN = os.path.join(BASE_DIR, "Formato Usuarios")
DESTINO = os.path.join(BASE_DIR, "plantillas")

# Caracter Wingdings que dibuja la casilla vacia en estos formatos
CASILLA = ""

# ── Etiquetas de celda → marcador. El numero es la ocurrencia: "Nombres"
#    aparece dos veces (funcionario y quien autoriza).
CAMPOS = [
    ("nombres", 1, "nombres"),
    ("apellidos", 1, "apellidos"),
    ("cedula de ciudadania", 1, "cedula"),
    ("telefono", 1, "telefono"),
    ("correo electronico", 1, "correo"),
    ("cargo", 1, "cargo"),
    ("oficina", 1, "oficina"),
    ("departamento", 1, "departamento"),
    ("ciudad", 1, "ciudad"),
    ("nombres", 2, "autNombres"),
    ("apellidos", 2, "autApellidos"),
    ("dependencia", 1, "autDependencia"),
    ("direccion ip del equipo donde funcionara", 1, "ani_ip"),
]

# ── Texto que acompaña a cada casilla → marcador ──────────────────────
CASILLAS = [
    # Tipo de solicitud
    ("creacion", "tipo_creacion"),
    ("actualizacion", "tipo_actualizacion"),
    # Vinculacion
    ("planta", "vinc_planta"),
    ("provisional", "vinc_provisional"),
    ("supernumerario", "vinc_supernumerario"),
    ("contratista", "vinc_contratista"),
    ("directivo", "vinc_directivo"),
    # Dependencia solicitante
    ("delegacion departamental", "dep_delegacion"),
    ("registraduria distrital", "dep_distrital"),
    ("registraduria especial/auxiliar/municipal", "dep_municipal"),
    ("oficinas centrales", "dep_centrales"),
    # Cargo de quien autoriza. Cada formato lo escribe distinto:
    # "Delegado(s) Departamental(es)" en unos, "Delegado Departamental"
    # en el SIRC; ambas variantes apuntan al mismo marcador.
    ("delegado(s)", "aut_delegado"),
    ("delegado departamental", "aut_delegado"),
    ("registrador(es)", "aut_registrador"),
    ("registrador distrital", "aut_registrador"),
    ("directivos", "aut_directivo"),
    ("coordinador de area", "aut_coordinador"),
    ("formador", "aut_formador"),
    # ── ANI: perfiles ──
    ("consultas ani", "p_consultas_ani"),
    ("archivo general", "p_archivo_general"),
    ("informacion ciudadana", "p_informacion_ciudadana"),
    ("recepcion secretaria", "p_recepcion_secretaria"),
    ("soporte tecnico interfaz", "p_soporte_tecnico"),
    ("centros de acopio", "p_centros_acopio"),
    ("derechos de peticion", "p_derechos_peticion"),
    ("notificacion resoluciones", "p_notificacion_resoluciones"),
    ("certificaciones vigencia", "p_certificaciones_vigencia"),
    ("correspondencia modificacion", "p_correspondencia"),
    ("registradores", "p_registradores"),
    ("duplicados pre envio", "p_duplicados"),
    ("cedulacion exterior", "p_cedulacion_exterior"),
    ("secretaria", "p_secretaria"),
    ("envios coordinacion", "p_envios_coordinacion"),
    # ── ANI: datos a consultar ──
    ("cedula de ciudadania", "d_cedula"),
    ("nombres y apellidos", "d_nombres"),
    ("lugar de expedicion", "d_lugar_expedicion"),
    ("fecha de expedicion", "d_fecha_expedicion"),
    ("vigencia", "d_vigencia"),
    ("resolucion", "d_resolucion"),
    ("fecha de nacimiento", "d_fecha_nacimiento"),
    ("lugar de nacimiento", "d_lugar_nacimiento"),
    ("grupo sanguineo", "d_grupo_sanguineo"),
    ("estatura", "d_estatura"),
    ("sexo", "d_sexo"),
    # ── Web Service ──
    ("sca", "a_sca"), ("ssc", "a_ssc"), ("cct", "a_cct"),
    ("inv. cc", "a_inv_cc"), ("inv. ti", "a_inv_ti"), ("verif", "a_verif"),
    ("1.1", "a_1_1"), ("1.n", "a_1_n"), ("cal", "a_cal"),
    ("impresion", "a_impresion"),
    ("reimpresion de documentos", "a_reimpresion"),
    # ── GED ──
    ("ged id", "h_ged_id"), ("ged rc", "h_ged_rc"),
    ("consulta", "a_consulta"),
    ("consulta e impresion", "a_consulta_impresion"),
    # ── SIRC ──
    ("estadisticas", "a_estadisticas"),
    ("consulta ged rcx", "a_consulta_ged"),
    ("correccion", "a_correccion"),
    ("expedicion certificado", "a_expedicion"),
    ("complementacion rcx", "a_complementacion"),
    ("correccion manual de anomalias", "a_correccion_manual"),
    ("borrado logico", "a_borrado"),
    ("dcu administrador", "a_dcu_admin"),
    ("dcu operador", "a_dcu_operador"),
    ("dcu registrador", "a_dcu_registrador"),
    ("grabacion rcx", "a_grabacion"),
    ("gestion de seriales (ans/gns)", "a_seriales"),
]


def limpiar(texto):
    """Minusculas, sin tildes, sin espacios extra ni dos puntos finales."""
    import unicodedata
    t = (texto or "").replace(CASILLA, "").strip()
    t = "".join(
        c for c in unicodedata.normalize("NFD", t)
        if unicodedata.category(c) != "Mn"
    )
    return " ".join(t.lower().replace(":", " ").split())


def forzar_arial(run, puntos=11):
    """Aplica Arial negro al run. Los runs nuevos nacen con la fuente por
    defecto de Word (Times New Roman) y los que sustituyen un texto guia
    ("dd /mm /aaaa") heredan su color gris: los datos van en negro."""
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    run.font.name = "Arial"
    run.font.size = Pt(puntos)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atributo), "Arial")


def poner_en_celda(celda, marcador):
    """Escribe {{marcador}} en la celda conservando su formato (y en Arial
    si la celda estaba vacia y no habia formato que heredar)."""
    parrafo = celda.paragraphs[0]
    if parrafo.runs:
        parrafo.runs[0].text = "{{" + marcador + "}}"
        forzar_arial(parrafo.runs[0])
        for r in parrafo.runs[1:]:
            r.text = ""
    else:
        forzar_arial(parrafo.add_run("{{" + marcador + "}}"))


def es_relleno(texto):
    """True si la celda solo trae adornos de plantilla (". . .", "dd/mm/aaaa",
    guiones bajos), es decir, sitio donde va un dato."""
    t = (texto or "").strip()
    if not t:
        return False
    return all(c in ". _-/dmaystDMAYST" for c in t)


def anexar_en_celda(celda, marcador):
    """Agrega el marcador al final del texto de la propia celda
    (para etiquetas que no tienen celda aparte, ej. "Teléfono:")."""
    parrafo = celda.paragraphs[0]
    if parrafo.runs:
        parrafo.runs[-1].text += "  {{" + marcador + "}}"
    else:
        forzar_arial(parrafo.add_run("  {{" + marcador + "}}"))


def marcar_campos(doc):
    """Recorre las celdas: cuando una contiene una etiqueta conocida, pone
    el marcador en la siguiente celda vacia de la misma fila.

    Una fila puede traer dos etiquetas (Cedula | ___ | Telefono | ___), por
    eso no se corta al encontrar la primera."""
    # etiqueta -> {ocurrencia: marcador}
    por_etiqueta = {}
    for etiqueta, ocurrencia, marcador in CAMPOS:
        por_etiqueta.setdefault(etiqueta, {})[ocurrencia] = marcador

    contador, puestos, vistas = {}, set(), set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = fila.cells
            for i, celda in enumerate(celdas):
                # Las celdas combinadas se repiten en fila.cells: saltarlas
                if celda._tc in vistas:
                    continue
                vistas.add(celda._tc)
                txt = limpiar(celda.text)
                if txt not in por_etiqueta:
                    continue
                contador[txt] = contador.get(txt, 0) + 1
                marcador = por_etiqueta[txt].get(contador[txt])
                if not marcador or marcador in puestos:
                    continue
                # 1) celda vacia a la derecha  2) celda con relleno (". . .")
                colocado = False
                for j in range(i + 1, len(celdas)):
                    if celdas[j]._tc is celda._tc:
                        continue
                    if not celdas[j].text.strip() or es_relleno(celdas[j].text):
                        poner_en_celda(celdas[j], marcador)
                        colocado = True
                        break
                # 3) sin celda disponible: anexar en la propia etiqueta
                if not colocado:
                    anexar_en_celda(celda, marcador)
                puestos.add(marcador)
    return puestos


def marcar_casillas(doc):
    """Reemplaza el caracter de casilla por {{cb:...}} identificando la
    opcion por el texto que la acompaña (mismo run, run siguiente o la
    celda vecina)."""
    puestas = set()
    mapa = dict(CASILLAS)

    def opcion_de(texto):
        """La opcion conocida mas larga con la que empieza el texto.
        Permite distinguir 'Consulta' de 'Consulta e Impresion'."""
        limpio = limpiar(texto)
        if not limpio:
            return None
        candidatas = [op for op in mapa if limpio.startswith(op)]
        return max(candidatas, key=len) if candidatas else None

    def procesar_parrafos(parrafos, texto_vecino=""):
        for parrafo in parrafos:
            runs = parrafo.runs
            for k, run in enumerate(runs):
                if CASILLA not in run.text:
                    continue
                # Texto que sigue a la casilla: lo que resta de su propio
                # run mas los runs siguientes, hasta la proxima casilla
                # (el texto puede estar a dos o tres runs de distancia)
                cola = run.text.split(CASILLA, 1)[1]
                for siguiente in runs[k + 1:]:
                    if CASILLA in siguiente.text:
                        cola += siguiente.text.split(CASILLA, 1)[0]
                        break
                    cola += siguiente.text
                opcion = opcion_de(cola) or opcion_de(texto_vecino)
                marcador = mapa.get(opcion) if opcion else None
                if not marcador or marcador in puestas:
                    continue
                run.text = run.text.replace(CASILLA, "{{cb:" + marcador + "}}", 1)
                puestas.add(marcador)

    # En las tablas "sí / no" cada opcion tiene DOS casillas en celdas
    # propias: la primera es "sí" y la segunda "no". Se numeran para que
    # el navegador pueda marcar la que corresponda en cada caso.
    veces = {}

    def marcar_celda_sola(celda, marcador):
        sufijo = "__no" if veces.get(marcador, 0) else ""
        veces[marcador] = veces.get(marcador, 0) + 1
        for parrafo in celda.paragraphs:
            for run in parrafo.runs:
                if CASILLA in run.text:
                    run.text = run.text.replace(
                        CASILLA, "{{cb:" + marcador + sufijo + "}}", 1
                    )
                    return True
        return False

    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = fila.cells
            # Orientacion de la fila: la mayoria de formatos ponen las
            # casillas primero y el texto despues ("❑ ❑ Consultas ANI"),
            # pero el GED lo hace al reves ("GED ID  ❑ ❑").
            texto_primero = bool(celdas) and bool(limpiar(celdas[0].text)) \
                and CASILLA not in celdas[0].text

            for i, celda in enumerate(celdas):
                # Etiqueta de la opcion, buscada en la direccion que
                # corresponda a la fila
                vecino = ""
                rango = (
                    range(i - 1, max(i - 4, -1), -1)
                    if texto_primero
                    else range(i + 1, min(i + 4, len(celdas)))
                )
                for j in rango:
                    t = celdas[j].text.strip()
                    # Saltar celdas de casilla, incluidas las que ya
                    # fueron reemplazadas por su marcador
                    if t and CASILLA not in t and "{{cb:" not in t:
                        vecino = t
                        break

                # Celda que contiene solo la casilla → tabla sí/no
                solo_casilla = CASILLA in celda.text and not limpiar(celda.text)
                if solo_casilla:
                    opcion = opcion_de(vecino)
                    marcador = mapa.get(opcion) if opcion else None
                    if marcador and marcar_celda_sola(celda, marcador):
                        puestas.add(marcador)
                    continue

                procesar_parrafos(celda.paragraphs, vecino)
    procesar_parrafos(doc.paragraphs)
    return puestas


def marcar_especiales(doc):
    """Campos que no siguen el patron etiqueta|celda-vacia:
    vigencia (Desde/Hasta en una misma celda), fecha de solicitud,
    justificacion y el nombre de usuario del tipo 'Actualizacion'."""
    puestos = set()

    def anexar_tras(parrafo, palabra, marcador):
        """Agrega el marcador justo despues de la palabra dentro del parrafo."""
        for run in parrafo.runs:
            if palabra.lower() in run.text.lower():
                run.text = run.text.rstrip() + "  {{" + marcador + "}}   "
                return True
        return False

    def marcar_vigencia(celda):
        """Reescribe la linea de vigencia como "Desde: … Hasta: …".

        Se reescribe entera (en vez de insertar los marcadores donde
        estaban) porque los formatos traen decenas de espacios entre
        ambas fechas y en las celdas angostas la linea se parte en dos."""
        TEXTO = "Desde: {{vigDesde}}  Hasta: {{vigHasta}}"
        for parrafo in celda.paragraphs:
            bajo = parrafo.text.lower()
            if "desde" not in bajo or "hasta" not in bajo:
                continue
            if parrafo.runs:
                parrafo.runs[0].text = TEXTO
                for r in parrafo.runs[1:]:
                    r.text = ""
            else:
                forzar_arial(parrafo.add_run(TEXTO))
            return True
        return False

    for tabla in doc.tables:
        filas = tabla.rows
        for idx, fila in enumerate(filas):
            vistas, celdas = set(), []
            for c in fila.cells:
                if c._tc not in vistas:
                    vistas.add(c._tc)
                    celdas.append(c)
            for i, celda in enumerate(celdas):
                txt = limpiar(celda.text)

                # Vigencia: "Desde: ____ Hasta: ____" dentro de una celda
                if "desde" in txt and "hasta" in txt and "vigDesde" not in puestos:
                    if marcar_vigencia(celda):
                        puestos.update({"vigDesde", "vigHasta"})


                # "Indique nombre de usuario______" se reemplaza entero por
                # un marcador: si el funcionario escribio un usuario va el
                # usuario; si no, vuelve a salir la frase guia.
                if "indique nombre de usuario" in txt and "textoUsuario" not in puestos:
                    for parrafo in celda.paragraphs:
                        runs = parrafo.runs
                        inicio = next(
                            (i for i, r in enumerate(runs) if "indique" in r.text.lower()),
                            None,
                        )
                        if inicio is None:
                            continue
                        # Se guarda el texto guia propio de este formato
                        # (unos traen guiones bajos y otros no) para que al
                        # ser "Creación" la linea quede igual al original.
                        # No se agregan espacios aqui: los pone el
                        # navegador solo cuando escribe un usuario.
                        guia = "".join(r.text for r in runs[inicio:]).strip()
                        guia = guia.replace("{", "").replace("}", "").replace("|", "")
                        runs[inicio].text = "{{textoUsuario|" + guia + "}}"
                        for r in runs[inicio + 1:]:
                            r.text = ""
                        puestos.add("textoUsuario")
                        break

                # Justificacion: la fila siguiente a "Describa la finalidad...".
                # En el GED esa frase comparte celda con el titulo
                # "JUSTIFICACIÓN DE ACCESO", por eso no basta con startswith.
                if "describa la finalidad" in txt and "justificacion" not in puestos:
                    if idx + 1 < len(filas):
                        destino = filas[idx + 1].cells[0]
                        if not destino.text.strip():
                            poner_en_celda(destino, "justificacion")
                            puestos.add("justificacion")
                    if "justificacion" not in puestos:
                        p = celda.add_paragraph()
                        forzar_arial(p.add_run("{{justificacion}}"))
                        puestos.add("justificacion")

    # Fecha de solicitud: la celda cuyo unico contenido es "dd /mm /aaaa"
    if "fechaSolicitud" not in puestos:
        vistas = set()
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda._tc in vistas:
                        continue
                    vistas.add(celda._tc)
                    t = celda.text.strip()
                    # El texto guia de la fecha es "dd /mm /aaaa" o
                    # "dd /mm /yyyy" segun el formato
                    tiene_ano = any(c in t.lower() for c in ("a", "y"))
                    if t and tiene_ano and es_relleno(t) and "/" in t:
                        poner_en_celda(celda, "fechaSolicitud")
                        puestos.add("fechaSolicitud")
                        return puestos
    return puestos


def uniformar_linea_actualizacion(doc):
    """La celda de 'Actualización · Indique nombre de usuario' mezcla dos
    tamaños de letra en el formato original. Se deja toda en Arial Narrow 9
    para que la linea se vea pareja (la casilla conserva Wingdings)."""
    from docx.oxml.ns import qn
    vistas = set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda._tc in vistas:
                    continue
                vistas.add(celda._tc)
                if "indique nombre de usuario" not in limpiar(celda.text):
                    continue
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        rpr = run._element.get_or_add_rPr()
                        rfonts = rpr.get_or_add_rFonts()
                        # El run de la casilla mantiene Wingdings: si le
                        # cambiamos la fuente, el recuadro desaparece
                        es_casilla = (
                            CASILLA in run.text or "{{cb:" in run.text
                        )
                        if es_casilla:
                            continue  # el recuadro conserva fuente y tamaño
                        run.font.name = "Arial Narrow"
                        for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                            rfonts.set(qn(atributo), "Arial Narrow")
                        run.font.size = Pt(9)
                return True
    return False


def unificar_arial(doc):
    """Pasada final: todo run que contenga un marcador de dato queda en
    Arial. Los marcadores de casilla ({{cb:...}}) se saltan porque viven
    en runs con fuente Wingdings, que es la que dibuja el recuadro."""
    def revisar(parrafos):
        for parrafo in parrafos:
            for run in parrafo.runs:
                if "{{" in run.text and "{{cb:" not in run.text:
                    forzar_arial(run)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                revisar(celda.paragraphs)
    revisar(doc.paragraphs)


def insertar_firmas(doc, ph_funcionario, ph_autorizador):
    """Coloca DOS espacios de firma: el primero para el funcionario y el
    segundo para quien autoriza, en el orden en que aparecen las etiquetas
    'Firma:' del formato.

    Cada firma va en la celda vecina si esta libre; si no, en la propia
    celda de la etiqueta (asi la firma del funcionario nunca se desplaza
    hasta la seccion de quien autoriza)."""
    ubicaciones, vistas = [], set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = []
            for c in fila.cells:
                if c._tc not in vistas:
                    vistas.add(c._tc)
                    celdas.append(c)
            for i, celda in enumerate(celdas):
                lim = limpiar(celda.text)
                # "Firma:" en su propia celda (formatos CDFT) o un bloque
                # que termina en "FIRMA" bajo una linea (formato GIFT05)
                if lim != "firma" and not lim.endswith("firma"):
                    continue
                destino = celda
                for j in range(i + 1, len(celdas)):
                    if not celdas[j].text.strip():
                        destino = celdas[j]
                        break
                ubicaciones.append(destino)

    puestas = 0
    for destino, imagen in zip(ubicaciones, [ph_funcionario, ph_autorizador]):
        # Por defecto va en el primer parrafo de la celda. Si la celda es
        # un bloque largo que termina en "FIRMA" (formato de Correo), la
        # imagen debe ir justo sobre esa linea, no al principio.
        parrafo = destino.paragraphs[0]
        for i, p in enumerate(destino.paragraphs):
            if limpiar(p.text) == "firma" and i > 0:
                parrafo = destino.paragraphs[i - 1]
                break
        parrafo.add_run().add_picture(imagen, width=Cm(5.5), height=Cm(1.5))
        puestas += 1
    return puestas


# ── GIFT05 (Correo): etiqueta → marcador. Este formato no usa celdas
#    ni casillas Wingdings, sino renglones con guiones bajos y cuadros
#    dibujados, por eso necesita su propio tratamiento.
GIFT05_CAMPOS = [
    ("Cuenta de usuario (únicamente para cuentas de servicio):", "cuentaUsuario"),
    ("Yo (nombre (s) y apellidos completos):", "nombreCompleto"),
    ("Documento de identidad:", "cedula"),
    ("Cargo (según resolución de nombramiento):", "cargo"),
    ("Oficina y/o dependencia:", "oficina"),
    ("Regional (CAN / Delegaciones / Registraduría  / Otro):", "regional"),
    ("Ciudad:", "ciudad"),
    ("Teléfono / extensión:", "telefono"),
    ("Nombre (s) y apellidos Completos:", "autNombreCompleto"),
    ("Cargo:", "autCargo"),
    ("Oficina y/o Dependencia:", "autDependencia"),
]

# Cuadros dibujados del GIFT05, EN EL ORDEN EN QUE APARECEN DENTRO DEL
# documento, que no coincide con el orden en que se ven en el papel.
# El orden se comprobo numerando cada cuadro y revisando el PDF:
#   solicitud → 1=Cambio 2=Bloqueo 3=Eliminación 4=Actualización 5=Creación
#   vinculación → 6=Supernumerario 7=Planta 8=Provisional 9=Libre Nombr.
GIFT05_CASILLAS = [
    ("Cambio de Responsable", "g_cambio"),
    ("Bloqueo", "g_bloqueo"),
    ("Eliminación", "g_eliminacion"),
    ("Actualización de datos", "g_actualizacion"),
    ("Creación", "g_creacion"),
    ("Supernumerario:", "g_supernumerario"),
    ("Planta:", "g_planta"),
    ("Provisional:", "g_provisional"),
    ("Libre Nombramiento:", "g_libre"),
    ("Contratista:", "g_contratista"),
    ("Entes de Control:", "g_entes"),
]


def preparar_gift05_campos(doc):
    """Renglones de guiones bajos → {{campo}} en el formato de Correo.

    Se trabaja parrafo por parrafo (no sobre el XML crudo) porque Word
    parte las etiquetas en varios trozos y no se pueden buscar enteras."""
    puestos = []

    # Etiqueta seguida de su linea de guiones bajos
    reglas = [(re.compile(re.escape(e) + r"\s*_{3,}"), e, m)
              for e, m in GIFT05_CAMPOS]
    # Textos guia que no usan guiones
    otras = [
        (re.compile(r"Fecha:\s*DD/MM/AAAA"), "Fecha:  {{fechaSolicitud}}"),
        (re.compile(r"Número\s*_{3,}"), "Número {{resNumero}}"),
        (re.compile(r"Año\s*_{3,}"), "Año {{resAno}}"),
        (re.compile(r"desde:\s*DÍA MES AÑO"), "desde:  {{vigDesde}}"),
        (re.compile(r"hasta:\s*DÍA MES AÑO"), "hasta:  {{vigHasta}}"),
        (re.compile(r"Fecha final nombramiento\s*DÍA MES AÑO"),
         "Fecha final nombramiento  {{vigHasta}}"),
    ]

    def procesar(parrafo):
        original = parrafo.text
        if not original.strip():
            return
        texto = original
        usados = []
        for patron, etiqueta, marcador in reglas:
            if marcador in puestos:
                continue
            texto, n = patron.subn(etiqueta + " {{" + marcador + "}}", texto, count=1)
            if n:
                usados.append(marcador)
        for patron, destino in otras:
            texto, n = patron.subn(destino, texto, count=1)
            if n:
                usados.append(destino[:12])
        # "Contratista" y "Entes de Control" no traen cuadro dibujado en
        # el formato: su marca se escribe como texto al lado
        for etiqueta, marcador in (("Contratista:", "xContratista"),
                                   ("Entes de Control:", "xEntes")):
            if marcador in puestos or etiqueta not in texto:
                continue
            texto = texto.replace(etiqueta, etiqueta + " {{" + marcador + "}}", 1)
            usados.append(marcador)
        if texto == original:
            return
        # Reescribir el parrafo conservando el formato del primer trozo
        if parrafo.runs:
            parrafo.runs[0].text = texto
            for r in parrafo.runs[1:]:
                r.text = ""
        else:
            forzar_arial(parrafo.add_run(texto))
        puestos.extend(usados)

    for tabla in doc.tables:
        for fila in tabla.rows:
            vistas = set()
            for celda in fila.cells:
                if celda._tc in vistas:
                    continue
                vistas.add(celda._tc)
                for parrafo in celda.paragraphs:
                    procesar(parrafo)
    for parrafo in doc.paragraphs:
        procesar(parrafo)
    return puestos


def gift05_tamano_uniforme(doc, puntos=10):
    """Deja en un solo tamaño los bloques del GIFT05 que vienen mas
    grandes que el resto (vigencia de acceso y la segunda linea del
    tipo de vinculacion). Debe ejecutarse DESPUES de unificar_arial,
    que de lo contrario los devuelve a 11 puntos."""
    marcas = ("resolucion de nombramiento", "vigencia cuenta de usuario",
              "contratista", "entes de control")
    ajustados = 0
    vistas = set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda._tc in vistas:
                    continue
                vistas.add(celda._tc)
                lim = limpiar(celda.text)
                if not any(m in lim for m in marcas):
                    continue
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.font.size = Pt(puntos)
                        ajustados += 1
    return ajustados


def preparar_gift05_casillas(ruta_docx):
    """Los cuadros de verificacion del GIFT05 son formas dibujadas con una
    letra adentro. Se numeran en el orden en que aparecen y su letra pasa
    a ser {{cbx:opcion}} para que el navegador escriba la X."""
    zin = zipfile.ZipFile(ruta_docx)
    contenido = {n: zin.read(n) for n in zin.namelist()}
    zin.close()
    xml = contenido["word/document.xml"].decode("utf-8")

    # Cada cuadro aparece dos veces (version moderna y de respaldo);
    # se recorren en orden y se saltan los que son etiquetas SI/NO
    bloques = list(re.finditer(r"<w:txbxContent>.*?</w:txbxContent>", xml, re.S))
    opciones = [m for _, m in GIFT05_CASILLAS]
    asignados, puestos = {}, []
    indice = 0
    piezas, ultimo = [], 0

    for bloque in bloques:
        texto = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", bloque.group(0)))
        limpio = texto.strip().upper()
        # Los cuadros del buzon ya traen escrito SÍ / NO: se marcan con
        # su propio marcador para dejar visible solo el que aplique
        if limpio in ("SI", "SÍ", "NO"):
            opcion = "buzon_si" if limpio != "NO" else "buzon_no"
            piezas.append(xml[ultimo:bloque.start()])
            # El texto puede venir partido ("S" + "Í"): el marcador va en
            # el primer trozo y los demas se vacian para no duplicarlo.
            # Ademas viene en gris muy claro (7F7F7F), casi invisible:
            # se reemplaza por negro para que la marca se lea.
            formato_buzon = (
                '<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
                '<w:b/><w:color w:val="000000"/><w:sz w:val="18"/>'
                '<w:szCs w:val="18"/></w:rPr>'
            )
            marca = "{{cbx:" + opcion + "}}"
            nuevo = re.sub(
                r"<w:r[^>]*>(?:<w:rPr>.*?</w:rPr>)?(<w:t[^>]*>)[^<]*(</w:t>)</w:r>",
                "<w:r>" + formato_buzon + r"\1" + marca + r"\2</w:r>",
                bloque.group(0), count=1, flags=re.S)
            if marca not in nuevo:
                # Estructura inesperada: se marca solo el texto
                nuevo = re.sub(r"(<w:t[^>]*>)[^<]*(</w:t>)",
                               r"\1" + marca + r"\2", bloque.group(0), count=1)
            # Vaciar los trozos siguientes para que el texto no se duplique
            antes, despues = nuevo.split(marca, 1)
            despues = re.sub(r"(<w:t[^>]*>)[^<]*(</w:t>)", r"\1\2", despues)
            nuevo = antes + marca + despues
            piezas.append(nuevo)
            ultimo = bloque.end()
            if opcion not in puestos:
                puestos.append(opcion)
            continue
        # Los cuadros duplicados comparten posicion visual: se agrupan
        clave = limpio
        if indice >= len(opciones):
            break
        marcador = opciones[indice]
        piezas.append(xml[ultimo:bloque.start()])
        # Los cuadros miden 18x10 pt pero traen 14 pt de margen interno:
        # una letra normal no cabe y Word la recorta. Se usa letra
        # pequeña (8 pt) y mas adelante se anulan esos margenes.
        formato = (
            '<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
            '<w:b/><w:color w:val="000000"/><w:sz w:val="16"/>'
            '<w:szCs w:val="16"/></w:rPr>'
        )
        piezas.append(
            re.sub(r"<w:r>(?:<w:rPr>.*?</w:rPr>)?(<w:t[^>]*>)[^<]*(</w:t>)</w:r>",
                   "<w:r>" + formato + r"\1{{cbx:" + marcador + r"}}\2</w:r>",
                   bloque.group(0), count=1, flags=re.S)
        )
        ultimo = bloque.end()
        asignados.setdefault(marcador, 0)
        asignados[marcador] += 1
        # Cada opcion ocupa dos copias del mismo cuadro
        if asignados[marcador] == 2:
            puestos.append("cbx:" + marcador)
            indice += 1
    piezas.append(xml[ultimo:])
    xml = "".join(piezas)

    # Anular el margen interno de los cuadros para que la marca quepa.
    # Los valores ya vienen definidos, asi que hay que reemplazarlos:
    # el formato antiguo (VML) usa "inset" y el moderno usa lIns/tIns/…
    xml = re.sub(r'(<v:textbox[^>]*?)inset="[^"]*"', r'\1inset="0,0,0,0"', xml)
    xml = re.sub(r"<v:textbox(?![^>]*inset=)", '<v:textbox inset="0,0,0,0"', xml)
    for lado in ("lIns", "tIns", "rIns", "bIns"):
        xml = re.sub(rf'({lado}=)"[^"]*"', rf'\1"0"', xml)
    xml = re.sub(r"<(wps|a):bodyPr(?![^>]*lIns=)",
                 r'<\1:bodyPr lIns="0" tIns="0" rIns="0" bIns="0"', xml)

    contenido["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(ruta_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for nombre, data in contenido.items():
            zout.writestr(nombre, data)
    return puestos


def firmas_delante_del_texto(ruta_docx):
    """Convierte las imagenes de firma de "en linea con el texto" a
    "delante del texto" (imagen flotante).

    Asi la firma se puede mover libremente en Word y no empuja el
    contenido ni agranda la celda. Los logos institucionales (.jpeg) no
    se tocan: solo las firmas (.png y .gif)."""
    zin = zipfile.ZipFile(ruta_docx)
    contenido = {n: zin.read(n) for n in zin.namelist()}
    zin.close()

    rels = contenido["word/_rels/document.xml.rels"].decode("utf-8")
    ids_firma = re.findall(
        r'Id="([^"]+)"[^>]*Target="media/[^"]*\.(?:png|gif)"', rels
    )
    if not ids_firma:
        return 0

    xml = contenido["word/document.xml"].decode("utf-8")
    convertidas = 0

    for bloque in re.findall(r"<wp:inline\b.*?</wp:inline>", xml, re.S):
        if not any(f'r:embed="{i}"' in bloque for i in ids_firma):
            continue

        # Piezas que hay que reordenar segun el esquema de wp:anchor
        extent = re.search(r"<wp:extent\b[^/]*/>", bloque)
        docpr = re.search(r"<wp:docPr\b[^/]*/>", bloque)
        frame = re.search(r"<wp:cNvGraphicFramePr>.*?</wp:cNvGraphicFramePr>", bloque, re.S)
        grafico = re.search(r"<a:graphic\b.*?</a:graphic>", bloque, re.S)
        if not (extent and docpr and grafico):
            continue

        # La firma se centra horizontalmente y se sube un poco para que
        # descanse sobre la linea de firma en lugar de taparla
        anchor = (
            '<wp:anchor xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
            ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
            ' distT="0" distB="0" distL="0" distR="0" simplePos="0"'
            ' relativeHeight="251658240" behindDoc="0" locked="0"'
            ' layoutInCell="1" allowOverlap="1">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="column"><wp:align>center</wp:align></wp:positionH>'
            '<wp:positionV relativeFrom="paragraph"><wp:posOffset>-260000</wp:posOffset></wp:positionV>'
            + extent.group(0)
            + '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            "<wp:wrapNone/>"
            + docpr.group(0)
            + (frame.group(0) if frame else "")
            + grafico.group(0)
            + "</wp:anchor>"
        )
        xml = xml.replace(bloque, anchor)
        convertidas += 1

    contenido["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(ruta_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for nombre, data in contenido.items():
            zout.writestr(nombre, data)
    return convertidas


def crear_placeholders(png, gif):
    """Dos imagenes transparentes. Se guardan en formatos distintos para
    que el navegador sepa cual reemplazar: .png = firma del funcionario,
    .gif = firma de quien autoriza (los logos del formato son .jpeg)."""
    from PIL import Image
    Image.new("RGBA", (540, 130), (255, 255, 255, 0)).save(png)
    Image.new("RGB", (540, 130), (255, 255, 255)).save(gif)


def procesar(nombre_archivo):
    base = nombre_archivo.replace(".docx", "")
    codigo = base.split(" - ")[0].strip().lower()
    origen = os.path.join(ORIGEN, nombre_archivo)
    destino = os.path.join(DESTINO, codigo + ".docx")
    shutil.copy2(origen, destino)

    doc = docx.Document(destino)
    if codigo == "gift05":
        # El formato de Correo no usa celdas etiqueta/valor: va aparte
        campos = set(preparar_gift05_campos(doc))
        casillas = set()
    else:
        campos = marcar_campos(doc)
        campos |= marcar_especiales(doc)
        casillas = marcar_casillas(doc)
    unificar_arial(doc)
    # Van despues de unificar_arial para que su formato propio prevalezca
    uniformar_linea_actualizacion(doc)
    if codigo == "gift05":
        gift05_tamano_uniforme(doc, puntos=10)

    ph_png = os.path.join(DESTINO, "_ph.png")
    ph_gif = os.path.join(DESTINO, "_ph.gif")
    crear_placeholders(ph_png, ph_gif)
    firmas = insertar_firmas(doc, ph_png, ph_gif)
    doc.save(destino)
    os.remove(ph_png)
    os.remove(ph_gif)

    # Los cuadros dibujados del formato de Correo se marcan sobre el XML
    if codigo == "gift05":
        casillas.update(preparar_gift05_casillas(destino))

    # Las firmas quedan "delante del texto" para poder reubicarlas en Word
    flotantes = firmas_delante_del_texto(destino)

    print(f"✅ {codigo}.docx — {len(campos)} campos, {len(casillas)} casillas,"
          f" {firmas} firma(s) ({flotantes} delante del texto)")
    faltan = [m for _, _, m in CAMPOS if m not in campos]
    if faltan:
        print(f"     campos sin ubicar: {', '.join(faltan)}")


def main():
    os.makedirs(DESTINO, exist_ok=True)
    for archivo in sorted(os.listdir(ORIGEN)):
        if archivo.endswith(".docx") and not archivo.startswith("~"):
            procesar(archivo)
    print(f"\nPlantillas en: {DESTINO}")


if __name__ == "__main__":
    main()
